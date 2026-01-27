import os
import re
import base64
import mistune
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_table_border(table):
    """给Word表格添加边框（全边框，1pt黑色）"""
    tbl_pr = table._element.tblPr
    tbl_borders = OxmlElement('w:tblBorders')

    # 定义所有边框类型（上、下、左、右、内横、内竖）
    borders = ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']
    for border in borders:
        elem = OxmlElement(f'w:{border}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')  # 边框宽度（4=1pt）
        elem.set(qn('w:color'), '000000')  # 黑色
        tbl_borders.append(elem)

    tbl_pr.append(tbl_borders)


import re


def parse_html_table(lines, start_idx):
    """
    解析HTML表格标签（<table>/<tr>/<td>/<th>），支持多行拆分的标签，返回表格数据和下一行索引
    :param lines: 所有md行列表
    :param start_idx: HTML表格起始行索引（<table>行）
    :return: (table_data, next_idx)
    """
    table_data = []
    idx = start_idx
    in_table = True
    in_tr = False  # 标记是否进入<tr>行
    current_row = []  # 存储当前<tr>中的单元格数据

    # 合并整个<table>块的所有内容为一个字符串（解决标签分行问题）
    table_content = []
    while idx < len(lines) and in_table:
        line = lines[idx].strip()
        table_content.append(line)
        if '</table>' in line:
            in_table = False
            break
        idx += 1
    # 合并成一个字符串，移除换行和多余空格，方便正则匹配
    table_str = ' '.join(table_content).replace('\n', ' ').replace('\r', ' ')
    # 移动到<table>结束后的下一行
    next_idx = idx + 1

    # 第一步：提取所有<tr>块（包括跨行的<tr>）
    tr_pattern = r'<tr[^>]*>(.*?)</tr>'
    tr_blocks = re.findall(tr_pattern, table_str, re.DOTALL)

    # 第二步：遍历每个<tr>块，提取<th>或<td>
    for tr_block in tr_blocks:
        row_data = []
        # 先提取<th>（表头）
        th_matches = re.findall(r'<th[^>]*>(.*?)</th>', tr_block, re.DOTALL)
        if th_matches:
            for th in th_matches:
                # 移除所有HTML标签，保留纯文本（处理<b>/<i>/<span>等嵌套标签）
                th_text = re.sub(r'<[^>]+>', '', th).strip()
                # 处理空单元格（避免出现空字符串）
                row_data.append(th_text if th_text else '')
            table_data.append(row_data)
            continue

        # 再提取<td>（内容）
        td_matches = re.findall(r'<td[^>]*>(.*?)</td>', tr_block, re.DOTALL)
        if td_matches:
            for td in td_matches:
                td_text = re.sub(r'<[^>]+>', '', td).strip()
                row_data.append(td_text if td_text else '')
            table_data.append(row_data)

    # 处理边界情况：如果没有提取到任何行数据，返回空列表
    if not table_data:
        print("警告：未从HTML表格中提取到任何数据")

    return table_data, next_idx

def parse_md_table(lines, start_idx):
    """解析Markdown原生表格，返回表格数据和下一行索引"""
    table_data = []
    idx = start_idx

    # 1. 读取表头
    header_line = lines[idx].strip()
    headers = [col.strip() for col in re.split(r'\|+', header_line) if col.strip()]
    table_data.append(headers)
    idx += 1

    # 2. 跳过分隔线行
    idx += 1

    # 3. 读取内容行
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or '|' not in line or line.startswith(('#', '-', '*', '+', '![')):
            break
        row = [col.strip() for col in re.split(r'\|+', line) if col.strip()]
        table_data.append(row)
        idx += 1

    return table_data, idx


def parse_md_image(line, md_file_dir):
    """解析Markdown图片语法（支持本地图片/Base64图片）"""
    # 匹配Markdown图片语法：![alt](path) 或 ![alt](data:image/;base64,xxx)
    img_pattern = r'!\[(.*?)\]\((.*?)(?:\s+".*?")?\)'
    match = re.match(img_pattern, line.strip())
    if not match:
        return None, None, False

    alt_text = match.group(1)  # 图片说明文字
    img_src = match.group(2).strip()  # 图片源（路径/Base64编码）

    # 1. 处理Base64图片
    if img_src.startswith('data:image/') and ';base64,' in img_src:
        try:
            # 拆分Base64头和编码内容
            base64_head, base64_data = img_src.split(';base64,', 1)
            # 解码Base64为二进制数据
            img_bytes = base64.b64decode(base64_data)
            # 写入内存流（BytesIO）
            img_data = BytesIO(img_bytes)
            return img_data, alt_text, True
        except Exception as e:
            print(f"Base64图片解码失败：{e}")
            return None, alt_text, True

    # 2. 处理本地图片
    if not os.path.isabs(img_src):
        img_src = os.path.join(md_file_dir, img_src)
        img_src = os.path.normpath(img_src)

    # 检查本地文件是否存在
    if os.path.exists(img_src) and os.path.isfile(img_src):
        return img_src, alt_text, False
    else:
        print(f"警告：本地图片文件不存在 - {img_src}")
        return None, alt_text, False


def md_to_docx(md_file_path, docx_file_path=None):
    """
    完整的Markdown转Word（支持HTML表格+Base64图片+Markdown表格+文本）
    """
    if not docx_file_path:
        docx_file_path = os.path.splitext(md_file_path)[0] + ".docx"

    # 获取md文件所在目录（用于拼接图片相对路径）
    md_file_dir = os.path.dirname(os.path.abspath(md_file_path))

    # 读取Markdown内容
    with open(md_file_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    # 创建Word文档并配置基础样式
    doc = Document()

    # 正文样式
    normal_style = doc.styles['Normal']
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.first_line_indent = Inches(0.5)
    normal_style.paragraph_format.line_spacing = 1.5

    # 标题样式（1-3级）
    for level in range(1, 4):
        style_name = f"Header {level}"
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "黑体"
        style.font.size = Pt(18 - 2 * level)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Inches(0)

    # 逐行解析Markdown
    lines = md_content.split("\n")
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            doc.add_paragraph("")  # 空行
            idx += 1
            continue

        # 1. 处理图片（优先级最高）
        img_data, alt_text, is_base64 = parse_md_image(line, md_file_dir)
        if img_data or is_base64:  # 即使Base64解码失败也处理说明文字
            # 添加图片说明文字（居中）
            if alt_text:
                p_alt = doc.add_paragraph(alt_text)
                p_alt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p_alt.font.size = Pt(10)
                p_alt.font.color.rgb = RGBColor(102, 102, 102)  # 灰色

            # 插入图片到Word
            if img_data:
                try:
                    # 添加图片，设置宽度为15cm（适配A4页面）
                    doc.add_picture(img_data, width=Cm(15))
                    # 图片居中
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # 图片后加空行
                    doc.add_paragraph("")
                except Exception as e:
                    print(f"插入图片失败：{e}")
                    # 插入失败时添加文字提示
                    p_error = doc.add_paragraph(f"[图片无法显示：{alt_text}]")
                    p_error.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            else:
                # Base64解码失败/本地图片不存在
                p_error = doc.add_paragraph(f"[图片无法显示：{alt_text}]")
                p_error.font.color.rgb = RGBColor(255, 0, 0)
            idx += 1
            continue

        # 2. 处理HTML表格（优先级高于Markdown原生表格）
        if line.startswith('<table') or ('<table>' in line and '</table>' not in line):
            table_data, idx = parse_html_table(lines, idx)
            row_count = len(table_data)
            col_count = len(table_data[0]) if row_count > 0 else 0
            if row_count > 0 and col_count > 0:
                table = doc.add_table(rows=row_count, cols=col_count)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                # 填充表格内容
                for i in range(row_count):
                    row_cells = table.rows[i].cells
                    for j in range(min(col_count, len(table_data[i]))):
                        cell = row_cells[j]
                        cell.text = table_data[i][j]
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "宋体"
                                run.font.size = Pt(11)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 添加表格边框
                set_table_border(table)
                doc.add_paragraph("")
            continue

        # 3. 处理Markdown原生表格
        if '|' in line and idx + 1 < len(lines) and re.match(r'^\|?\s*:?-+:?\s*\|.*$', lines[idx + 1].strip()):
            table_data, idx = parse_md_table(lines, idx)
            row_count = len(table_data)
            col_count = len(table_data[0]) if row_count > 0 else 0
            if row_count > 0 and col_count > 0:
                table = doc.add_table(rows=row_count, cols=col_count)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                # 填充表格内容
                for i in range(row_count):
                    row_cells = table.rows[i].cells
                    for j in range(min(col_count, len(table_data[i]))):
                        cell = row_cells[j]
                        cell.text = table_data[i][j]
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "宋体"
                                run.font.size = Pt(11)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 添加表格边框
                set_table_border(table)
                doc.add_paragraph("")
            continue

        # 4. 处理标题
        if line.startswith("# "):
            doc.add_paragraph(line[2:], style="Header 1")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Header 2")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Header 3")
        # 5. 处理无序列表
        elif line.startswith(("- ", "* ", "+ ")):
            p = doc.add_paragraph(line[2:], style="Normal")
            p.paragraph_format.first_line_indent = Inches(0)
        # 6. 处理粗体/斜体和普通文本（过滤HTML标签）
        else:
            # 移除剩余的HTML标签，保留纯文本
            clean_text = re.sub(r'<[^>]+>', '', line)
            p = doc.add_paragraph(style="Normal")
            patterns = [
                (r'\*\*(.*?)\*\*', True, False),  # 粗体
                (r'\*(.*?)\*', False, True)  # 斜体
            ]
            pos = 0
            for pattern, is_bold, is_italic in patterns:
                for match in re.finditer(pattern, clean_text):
                    if match.start() > pos:
                        run = p.add_run(clean_text[pos:match.start()])
                        run.font.name = "宋体"
                        run.font.size = Pt(12)
                    run = p.add_run(match.group(1))
                    run.font.name = "宋体"
                    run.font.size = Pt(12)
                    run.bold = is_bold
                    run.italic = is_italic
                    pos = match.end()
            if pos < len(clean_text):
                run = p.add_run(clean_text[pos:])
                run.font.name = "宋体"
                run.font.size = Pt(12)

        idx += 1

    # 保存Word文件
    doc.save(docx_file_path)
    print(f"转换完成（含HTML表格+Base64图片）：{docx_file_path}")

def batch_md_to_docx(folder_path):
    """批量转换文件夹下所有md文件"""
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".md"):
            md_file = os.path.join(folder_path, file_name)
            md_to_docx(md_file)

# 运行示例
if __name__ == "__main__":
    # 单文件转换（替换为你的md文件路径）
    md_to_docx("/mnt/d/BaiduNetdiskDownload/样例pdf/markdown/1957_330783-1-1 东阳县人民委员会颁布“东阳县文物保护单位目录（第一批）”的通知（（57）东文字第1029号）1957年5月29日.md")

    # 批量转换（替换为你的文件夹路径）
    # batch_md_to_docx("你的文件夹路径")